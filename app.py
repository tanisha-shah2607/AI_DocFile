import os
from groq import Groq
from docx import Document

# API key 
GROQ_API_KEY = "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"

client = Groq(api_key=GROQ_API_KEY)

# text generation
def generate_ai_text(prompt):
    print("\nGenerating AI content using Groq...")

    response = client.chat.completions.create(
        model="llama-3.1-8b-instant", 
        messages=[
            {"role": "system", "content": "You are a helpful assistant that writes clear, structured content."},
            {"role": "user", "content": prompt}],temperature=0.7,max_tokens=2000)

    return response.choices[0].message.content

# doc file creation
def create_docx(content, filename="AI_Document.docx"):
    doc = Document()
    doc.add_heading("AI Generated Content", level=1)
    doc.add_paragraph(content)
    
    # Saving in the vs code
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    
    print(f"\nDocument saved as: {filepath}")
    print(f"Location: {os.getcwd()}")

